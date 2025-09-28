import os
import requests
import streamlit as st
import chromadb
from contract_ingest import load_contract, split_into_clauses, GTEEmbeddingFunction
from docx import Document  # ⬅️ 用來生成 Word
from io import BytesIO
import io
from docx import Document
import json
from datetime import datetime

# ==========================
# Ollama 設定
# ==========================
OLLAMA_API_URL = "http://localhost:11434/api/generate"

AVAILABLE_MODELS = [
    "qwen3:4b",
    "qwen3:8b",
    "chatglm3:6b",
    "llama2:13b-chat",
    "mistral:7b-instruct"
]

def call_ollama(prompt, model_name, max_tokens=500):
    payload = {
        "model": model_name,
        "prompt": prompt,
        "options": {"temperature": 0.7},
        "stream": False
    }
    resp = requests.post(OLLAMA_API_URL, json=payload)
    if resp.status_code == 200:
        return resp.json().get("response", "").strip()
    else:
        return f"❌ Ollama 請求失敗: {resp.text}"

# ==========================
# 初始化向量資料庫
# ==========================
client = chromadb.PersistentClient(path="./chroma_db")
laws_collection = client.get_collection(
    name="hk_cap4_laws", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)
contracts_collection = client.get_or_create_collection(
    name="contracts", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)

# ==========================
# Streamlit UI
# ==========================
st.set_page_config(page_title="合約分析系統", layout="wide")
st.title("📑 本地合約分析系統（Ollama + RAG）")

selected_model = st.sidebar.selectbox("選擇 Ollama 模型", AVAILABLE_MODELS, index=0)

uploaded_file = st.file_uploader("上傳合約 (PDF / DOCX / TXT)", type=["pdf", "docx", "txt"])

if uploaded_file:
    # 1. 解析文件
    with open(f"./contracts/{uploaded_file.name}", "wb") as f:
        f.write(uploaded_file.getbuffer())
    text = load_contract(f"./contracts/{uploaded_file.name}")
    clauses = split_into_clauses(text, max_len=600)

    st.success(f"✅ 成功載入合約，共 {len(clauses)} 個條款")

    tab1, tab2, tab3 = st.tabs(["📖 條款逐條分析", "📌 合約摘要", "⚠️ 風險重點"])

    clause_analyses = []

    # 條款逐條分析
    with tab1:
        for i, clause in enumerate(clauses):
            with st.spinner(f"分析第 {i+1} 條款中..."):
                prompt = f"""
請閱讀以下合約條款，並分析：
1. 條款摘要
2. 潛在風險與爭議
3. 可能涉及的法律依據

條款內容：
{clause}
"""
                analysis = call_ollama(prompt, selected_model, max_tokens=500)
                clause_analyses.append((clause, analysis))
                st.markdown(f"### 條款 {i+1}")
                st.info(clause)
                st.write(analysis)

    # 合約摘要
    with tab2:
        with st.spinner("正在生成合約摘要..."):
            summary_prompt = f"請閱讀以下合約全文，生成一份簡潔的摘要：\n\n{text[:6000]}"
            summary = call_ollama(summary_prompt, selected_model, max_tokens=600)
        st.subheader("📌 合約摘要")
        st.write(summary)

    # 風險重點
    with tab3:
        with st.spinner("正在生成風險重點..."):
            risk_prompt = f"請閱讀以下合約全文，列出可能存在的風險與爭議條款：\n\n{text[:6000]}"
            risks = call_ollama(risk_prompt, selected_model, max_tokens=600)
        st.subheader("⚠️ 風險重點")
        st.write(risks)

# ====== 生成 Word 報告函式 ======
def generate_word_report(summary, risks, clause_analyses):
    doc = Document()
    doc.add_heading("合約分析報告", level=1)

    # 摘要
    doc.add_heading("📌 合約摘要", level=2)
    doc.add_paragraph(summary if summary else "無摘要")

    # 潛在風險
    doc.add_heading("⚠️ 潛在風險", level=2)
    if risks:
        if isinstance(risks, str):  # 確保是 list
            risks = risks.split("\n")
        for r in risks:
            doc.add_paragraph(f"- {r.strip()}")
    else:
        doc.add_paragraph("未檢測到明顯風險。")

    # 條款分析
    doc.add_heading("📑 條款逐條分析", level=2)
    for i, (clause, analysis) in enumerate(clause_analyses, 1):
        doc.add_heading(f"條款 {i}", level=3)
        doc.add_paragraph(clause, style="Intense Quote")
        doc.add_paragraph(f"🔎 分析: {analysis}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ====== Streamlit UI 裡加下載按鈕 ======
if uploaded_file and st.button("📥 生成並下載 Word 報告"):
    buf = generate_word_report(summary, risks, clause_analyses)
    st.download_button(
        label="⬇️ 下載合約分析報告 (Word)",
        data=buf,
        file_name="contract_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ====== 生成 JSON 記錄 ======
def save_json_report(summary, risks, clause_analyses, filename="contract_analysis.json"):
    data = {
        "generated_at": datetime.now().isoformat(),
        "summary": summary,
        "risks": risks if isinstance(risks, list) else risks.split("\n"),
        "clauses": [
            {"clause": clause, "analysis": analysis}
            for clause, analysis in clause_analyses
        ]
    }

    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return filename


# ====== Streamlit UI 新增保存按鈕 ======
if uploaded_file and st.button("📥 生成並下載 JSON + Word 報告"):
    # 生成 Word
    buf = generate_word_report(summary, risks, clause_analyses)
    st.download_button(
        label="⬇️ 下載合約分析報告 (Word)",
        data=buf,
        file_name="contract_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # 自動保存 JSON
    json_file = save_json_report(summary, risks, clause_analyses)
    st.success(f"✅ JSON 記錄已保存到本地：{json_file}")