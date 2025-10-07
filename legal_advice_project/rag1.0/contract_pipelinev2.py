import os
import re
import json
from datetime import datetime
from docx import Document
from dotenv import load_dotenv

from contract_ingest import load_contract, split_into_clauses
from vertexai import rag
from vertexai.generative_models import Tool, GenerativeModel

# ==========================
# 環境變數
# ==========================
load_dotenv()
PROJECT_ID = os.getenv("GCP_PROJECT")
LOCATION = os.getenv("GCP_LOCATION", "us-central1")
RAG_CORPUS_NAME = os.getenv("RAG_CORPUS_NAME")  # 你要在 .env 設定這個值

if not PROJECT_ID or not RAG_CORPUS_NAME:
    raise ValueError("❌ 請在 .env 設定 GCP_PROJECT 與 RAG_CORPUS_NAME")

REPORTS_DIR = "./reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# ==========================
# 初始化 Vertex AI RAG
# ==========================
rag_retrieval_config = rag.RagRetrievalConfig(
    top_k=3,
    filter=rag.Filter(vector_distance_threshold=0.5)
)

rag_tool = Tool.from_retrieval(
    retrieval=rag.Retrieval(
        source=rag.VertexRagStore(
            rag_resources=[rag.RagResource(rag_corpus=RAG_CORPUS_NAME)],
            rag_retrieval_config=rag_retrieval_config
        )
    )
)

# 第一層 Gemini（生成初稿）
rag_model_primary = GenerativeModel(
    model_name="gemini-2.0-flash-001",
    tools=[rag_tool]
)

# 第二層 Gemini（複核答案）
rag_model_reviewer = GenerativeModel(
    model_name="gemini-2.0-flash-001"
)

# ==========================
# 清理輸出，移除冗餘字眼
# ==========================
def clean_output(text: str) -> str:
    if not text:
        return ""
    # 移除常見開頭
    text = re.sub(r"^以下是.*\n?", "", text)
    text = re.sub(r"^好的，我來.*\n?", "", text)
    # 移除「修正說明」、「修正後」等字眼
    text = text.replace("修正說明及理由：", "")
    text = text.replace("修正後的風險分析：", "")
    text = text.replace("修正後的摘要：", "")
    text = text.replace("修正後的分析：", "")
    return text.strip()

# ==========================
# 條款分析（使用 RAG + 複核）
# ==========================
def analyze_clause(clause_text: str):
    try:
        # 初稿
        prompt_primary = f"""
請分析以下合約條款：

條款內容：
{clause_text[:800]}

請提供：
1. 條款要點
2. 潛在風險
3. 法律依據
"""
        draft = rag_model_primary.generate_content(prompt_primary).text.strip()

        # 複核（要求乾淨輸出）
        prompt_review = f"""
你是一位嚴謹的法律審核助手。請直接輸出最終分析內容，不要包含「以下是」、「修正說明」、「修正後」等字眼，也不要描述審核過程。

條款內容：
{clause_text[:800]}

初步分析：
{draft}

請輸出乾淨、正式的最終分析：
"""
        final = rag_model_reviewer.generate_content(prompt_review).text.strip()
        return clean_output(final)

    except Exception as e:
        print(f"❌ 條款分析錯誤: {e}")
        return f"條款分析失敗: {str(e)}"

# ==========================
# 全局合約分析（使用 RAG + 複核）
# ==========================
def analyze_contract_global(contract_text: str):
    limited_text = contract_text[:6000]

    summary_prompt = f"""
請分析以下合約內容，提供簡潔的摘要：
1. 合約類型和目的
2. 主要當事人
3. 核心條款要點

合約內容：
{limited_text}
"""
    risk_prompt = f"""
請分析以下合約內容，識別潛在風險：
1. 法律風險
2. 商業風險
3. 執行風險

合約內容：
{limited_text}
"""

    print("🔍 生成合約摘要...")
    draft_summary = rag_model_primary.generate_content(summary_prompt).text.strip()
    summary_review_prompt = f"""
你是一位嚴謹的法律審核助手。請直接輸出最終摘要，不要包含「以下是」、「修正說明」、「修正後」等字眼。

合約內容：
{limited_text[:1000]}...

初稿摘要：
{draft_summary}

請輸出乾淨、正式的最終摘要：
"""
    summary = rag_model_reviewer.generate_content(summary_review_prompt).text.strip()
    summary = clean_output(summary)

    print("⚠️ 分析潛在風險...")
    draft_risks = rag_model_primary.generate_content(risk_prompt).text.strip()
    risks_review_prompt = f"""
你是一位嚴謹的法律審核助手。請直接輸出最終風險分析，不要包含「以下是」、「修正說明」、「修正後」等字眼。

合約內容：
{limited_text[:1000]}...

初稿風險分析：
{draft_risks}

請輸出乾淨、正式的最終風險分析：
"""
    risks = rag_model_reviewer.generate_content(risks_review_prompt).text.strip()
    risks = clean_output(risks)

    return {"summary": summary, "risks": risks}

# ==========================
# 報告輸出
# ==========================
def generate_word_report(filename, summary, risks, clause_analyses):
    doc = Document()
    doc.add_heading("合約分析報告", level=1)

    doc.add_heading("📌 合約摘要", level=2)
    doc.add_paragraph(summary if summary else "無摘要")

    doc.add_heading("⚠️ 潛在風險", level=2)
    if risks:
        for r in risks.split("\n"):
            doc.add_paragraph(f"- {r.strip()}")
    else:
        doc.add_paragraph("未檢測到明顯風險。")

    doc.add_heading("📑 條款逐條分析", level=2)
    for i, (clause, analysis) in enumerate(clause_analyses, 1):
        doc.add_heading(f"條款 {i}", level=3)
        doc.add_paragraph(clause, style="Intense Quote")
        doc.add_paragraph(f"🔎 分析: {analysis}")

    base_name = os.path.basename(filename)
    out_name = base_name.rsplit(".", 1)[0] + "_分析.docx"
    out_path = os.path.join(REPORTS_DIR, out_name)
    doc.save(out_path)
    return out_path

def save_json_report(filename, summary, risks, clause_analyses):
    data = {
        "generated_at": datetime.now().isoformat(),
        "summary": summary,
        "risks": risks.split("\n") if isinstance(risks, str) else risks,
        "clauses": [{"clause": c, "analysis": a} for c, a in clause_analyses]
    }

    base_name = os.path.basename(filename)
    out_name = base_name.rsplit(".", 1)[0] + "_分析.json"
    out_path = os.path.join(REPORTS_DIR, out_name)

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return out_path

# ==========================
# 主流程：掃描 contracts/
# ==========================
def analyze_contract_file(file_path):
    print(f"\n🚀 開始分析合約：{file_path}")
    text = load_contract(file_path)
    clauses = split_into_clauses(text, max_len=600)

    clause_analyses = []
    for i, clause in enumerate(clauses, 1):
        print(f"🔎 條款 {i} 分析中...")
        analysis = analyze_clause(clause)
        clause_analyses.append((clause, analysis))

    global_analysis = analyze_contract_global(text)
    summary, risks = global_analysis["summary"], global_analysis["risks"]

    word_file = generate_word_report(file_path, summary, risks, clause_analyses)
    json_file = save_json_report(file_path, summary, risks, clause_analyses)

    print(f"✅ 報告完成：{word_file}, {json_file}")
    return {
        "word": word_file,
        "json": json_file,
        "summary": summary,
        "risks": risks,
        "clauses": clause_analyses
    }

if __name__ == "__main__":
    contracts_dir = "./contracts"
    for filename in os.listdir(contracts_dir):
        file_path = os.path.join(contracts_dir, filename)
        if os.path.isfile(file_path) and filename.lower().endswith((".pdf", ".docx", ".txt")):
            analyze_contract_file(file_path)
