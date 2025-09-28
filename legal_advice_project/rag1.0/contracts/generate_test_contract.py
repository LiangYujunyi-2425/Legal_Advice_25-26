from docx import Document

doc = Document()

doc.add_heading("僱傭合約", level=1)
doc.add_paragraph("本合約由以下雙方訂立：")
doc.add_paragraph("僱主：ABC 公司")
doc.add_paragraph("僱員：張三")

doc.add_paragraph("第一條（工作內容）")
doc.add_paragraph("僱員需履行公司安排之職務，包括但不限於文件處理、客戶服務及資料輸入。")

doc.add_paragraph("第二條（工作時間）")
doc.add_paragraph("僱員每週工作 48 小時，星期一至星期六，每日 8 小時。")

doc.add_paragraph("第三條（薪酬）")
doc.add_paragraph("僱員之月薪為港幣 10,000 元，每月最後一日支付。")

doc.add_paragraph("第四條（休息日及假期）")
doc.add_paragraph("僱員每週享有一天休息日，並享有法定假期。")

doc.add_paragraph("第五條（合約期限）")
doc.add_paragraph("本合約期限為兩年，自 2025 年 1 月 1 日起至 2026 年 12 月 31 日止。")

doc.add_paragraph("第六條（解約通知）")
doc.add_paragraph("任何一方欲終止合約，須提前 3 日以書面通知對方。")

doc.add_paragraph("第七條（競業限制）")
doc.add_paragraph("僱員於離職後 12 個月內，不得在香港任何與公司業務相關之企業任職。")

doc.add_paragraph("第八條（保密條款）")
doc.add_paragraph("僱員不得於任職期間及離職後，向第三方披露有關公司之商業秘密。")

doc.add_paragraph("第九條（其他）")
doc.add_paragraph("本合約受香港特別行政區法律管轄。")

doc.save("employment_contract_test.docx")
print("✅ 已生成 employment_contract_test.docx")
