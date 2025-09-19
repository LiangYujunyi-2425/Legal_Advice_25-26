import os
import json
import requests
import chromadb
from io import BytesIO
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
from contract_ingest import load_contract, split_into_clauses, GTEEmbeddingFunction

# ==========================
# 環境變數
# ==========================
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("❌ 請先在 .env 設定 GEMINI_API_KEY")

GEMINI_MODEL = "gemini-1.5-flash"
GEMINI_ENDPOINT = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"
HEADERS = {"Content-Type": "application/json", "x-goog-api-key": GEMINI_API_KEY}

REPORTS_DIR = "./reports"
os.makedirs(REPORTS_DIR, exist_ok=True)  # 如果沒有資料夾就建立

# ==========================
# 初始化向量資料庫
# ==========================
client_chroma = chromadb.PersistentClient(path="./chroma_db")
laws_collection = client_chroma.get_collection(
    name="hk_cap4_laws", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)
contracts_collection = client_chroma.get_or_create_collection(
    name="contracts", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)

# ==========================
# Gemini API 呼叫
# ==========================
def call_gemini(prompt: str, temperature=0.6, max_tokens=800):
    body = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens}
    }
    resp = requests.post(GEMINI_ENDPOINT, headers=HEADERS, json=body)
    if resp.status_code != 200:
        return f"⚠️ Gemini 請求失敗: {resp.status_code} {resp.text}"
    try:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
    except Exception:
        return str(resp.json())

# ==========================
# 雙 Gemini 流程
# ==========================
def dual_gemini_answer(prompt: str):
    draft = call_gemini(prompt, temperature=0.7, max_tokens=800)
    review_prompt = f"""
以下是第一個 Gemini 模型生成的草稿，請你檢查是否完整涵蓋重點、是否有錯誤或幻覺，
並在保持準確性的前提下，務必使用 **繁體中文** 重新輸出最終回答。

草稿回答：
{draft}
"""
    final = call_gemini(review_prompt, temperature=0.3, max_tokens=800)
    return final

# ==========================
# 條款分析
# ==========================
def analyze_clause(clause_text: str):
    law_results = laws_collection.query(query_texts=[clause_text], n_results=3)
    law_refs = [doc for doc in law_results["documents"][0]]
    contract_results = contracts_collection.query(query_texts=[clause_text], n_results=2)
    contract_refs = [doc for doc in contract_results["documents"][0]]

    context = "\n".join(law_refs + contract_refs)

    prompt = f"""
請閱讀以下合約條款，並分析：
1. 條款摘要
2. 潛在風險與爭議
3. 可能涉及的法律依據

條款內容：
{clause_text}

相關法律與參考：
{context}
"""
    return dual_gemini_answer(prompt)

# ==========================
# 全局合約分析
# ==========================
def analyze_contract_global(contract_text: str):
    summary_prompt = f"請閱讀以下合約全文，生成一份摘要，包含合約的主要內容與目的：\n{contract_text[:8000]}"
    risk_prompt = f"請閱讀以下合約全文，列出可能存在的風險與爭議條款：\n{contract_text[:8000]}"

    summary = dual_gemini_answer(summary_prompt)
    risks = dual_gemini_answer(risk_prompt)
    return {"summary": summary, "risks": risks}

# ==========================
# 報告輸出
# ==========================
def generate_word_report(filename, summary, risks, clause_analyses):
    doc = Document()
    doc.add_heading("合約分析報告", level=1)

    doc.add_heading("📌 合約摘要", level=2)
    doc.add_paragraph(summary if summary else "無摘要")

    doc.add_heading("⚠️ 潛在風險", level=2)
    if risks:
        for r in risks.split("\n"):
            doc.add_paragraph(f"- {r.strip()}")
    else:
        doc.add_paragraph("未檢測到明顯風險。")

    doc.add_heading("📑 條款逐條分析", level=2)
    for i, (clause, analysis) in enumerate(clause_analyses, 1):
        doc.add_heading(f"條款 {i}", level=3)
        doc.add_paragraph(clause, style="Intense Quote")
        doc.add_paragraph(f"🔎 分析: {analysis}")

    base_name = os.path.basename(filename)
    out_name = base_name.rsplit(".", 1)[0] + "_分析.docx"
    out_path = os.path.join(REPORTS_DIR, out_name)
    doc.save(out_path)
    return out_path


def save_json_report(filename, summary, risks, clause_analyses):
    data = {
        "generated_at": datetime.now().isoformat(),
        "summary": summary,
        "risks": risks.split("\n") if isinstance(risks, str) else risks,
        "clauses": [{"clause": c, "analysis": a} for c, a in clause_analyses]
    }

    base_name = os.path.basename(filename)
    out_name = base_name.rsplit(".", 1)[0] + "_分析.json"
    out_path = os.path.join(REPORTS_DIR, out_name)

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return out_path

# ==========================
# 主流程：掃描 contracts/
# ==========================
def analyze_contract_file(file_path):
    print(f"\n🚀 開始分析合約：{file_path}")
    text = load_contract(file_path)
    clauses = split_into_clauses(text, max_len=600)

    clause_analyses = []
    for i, clause in enumerate(clauses, 1):
        print(f"🔎 條款 {i} 分析中...")
        analysis = analyze_clause(clause)
        clause_analyses.append((clause, analysis))

    global_analysis = analyze_contract_global(text)
    summary, risks = global_analysis["summary"], global_analysis["risks"]

    word_file = generate_word_report(file_path, summary, risks, clause_analyses)
    json_file = save_json_report(file_path, summary, risks, clause_analyses)

    print(f"✅ 報告完成：{word_file}, {json_file}")

if __name__ == "__main__":
    contracts_dir = "./contracts"
    for filename in os.listdir(contracts_dir):
        file_path = os.path.join(contracts_dir, filename)
        if os.path.isfile(file_path) and filename.lower().endswith((".pdf", ".docx", ".txt")):
            analyze_contract_file(file_path)
