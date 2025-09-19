import os
import requests
import streamlit as st
import chromadb
from contract_ingest import load_contract, split_into_clauses, GTEEmbeddingFunction
from docx import Document
from io import BytesIO
import json
from datetime import datetime
import pickle
from rank_bm25 import BM25Okapi
import jieba
from sentence_transformers import SentenceTransformer, CrossEncoder
import argparse

parser = argparse.ArgumentParser()
parser.add_argument("--generator", type=str, default="mistral:7b-instruct")
parser.add_argument("--verifier", type=str, default="qwen3:8b")
args = parser.parse_args()

GENERATOR_MODEL = args.generator
VERIFIER_MODEL = args.verifier

# ==========================
# Ollama 設定
# ==========================
OLLAMA_API_URL = "http://localhost:11434/api/generate"

def call_ollama(model_name, prompt, max_tokens=500):
    payload = {
        "model": model_name,
        "prompt": prompt,
        "options": {"temperature": 0.3},
        "stream": False
    }
    resp = requests.post(OLLAMA_API_URL, json=payload)
    if resp.status_code == 200:
        return resp.json().get("response", "").strip()
    else:
        return f"❌ Ollama 請求失敗: {resp.text}"

# ==========================
# 初始化向量資料庫
# ==========================
client = chromadb.PersistentClient(path="./chroma_db")
laws_collection = client.get_collection(
    name="hk_cap4_laws", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)
contracts_collection = client.get_or_create_collection(
    name="contracts", embedding_function=GTEEmbeddingFunction("thenlper/gte-large-zh")
)

# ==========================
# BM25 (載入索引)
# ==========================
with open("bm25_index.pkl", "rb") as f:
    bm25_data = pickle.load(f)
bm25 = bm25_data["bm25"]
bm25_chunks = bm25_data["chunks"]

embedder = SentenceTransformer("thenlper/gte-large-zh")
reranker = CrossEncoder("BAAI/bge-reranker-large")

# ==========================
# Hybrid Search
# ==========================
def hybrid_search(query: str, n=10):
    vector_results = laws_collection.query(
        query_texts=[query],
        n_results=n,
        include=["documents", "metadatas", "distances"]
    )
    vector_docs = vector_results["documents"][0]
    vector_metas = vector_results["metadatas"][0]
    vector_scores = vector_results["distances"][0]

    vector_candidates = [
        (doc, meta, 1 - score, "Chroma")
        for doc, meta, score in zip(vector_docs, vector_metas, vector_scores)
    ]

    tokenized_query = list(jieba.cut(query))
    bm25_scores = bm25.get_scores(tokenized_query)
    top_idx = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:n]

    bm25_candidates = [
        (bm25_chunks[i]["text"], bm25_chunks[i], bm25_scores[i], "BM25") for i in top_idx
    ]

    merged = {}
    for doc, meta, score, source in vector_candidates + bm25_candidates:
        if doc not in merged or score > merged[doc][1]:
            merged[doc] = (meta, score, source)

    return [(doc, meta, score, source) for doc, (meta, score, source) in merged.items()]

def rerank(query, candidates, top_k=3):
    pairs = [(query, doc) for doc, _, _, _ in candidates]
    scores = reranker.predict(pairs)
    ranked = sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)
    return ranked[:top_k]

# ==========================
# 雙 LLM Pipeline
# ==========================
def generate_answer(query, context_texts):
    prompt = (
        "你是一位香港法律輔助助手，請根據以下條文生成初步回答：\n\n"
        f"條文：\n{chr(10).join(context_texts)}\n\n問題：{query}"
    )
    return call_ollama(GENERATOR_MODEL, prompt, max_tokens=512)

def verify_answer(query, draft_answer, context_texts):
    prompt = (
        "你是一位嚴謹的法律審核助手。以下是初步回答與法律條文：\n\n"
        f"問題：{query}\n\n"
        f"初步回答：{draft_answer}\n\n"
        f"條文：\n{chr(10).join(context_texts)}\n\n"
        "請檢查並修正錯誤，補充遺漏，並加強引用，保持繁體中文。"
    )
    return call_ollama(VERIFIER_MODEL, prompt, max_tokens=700)

# ==========================
# Streamlit UI
# ==========================
st.set_page_config(page_title="合約 + 法律分析系統", layout="wide")
st.title("📑 本地合約 + 法律分析系統（Ollama + 雙 LLM + RAG）")

uploaded_file = st.file_uploader("📂 上傳合約 (PDF / DOCX / TXT)", type=["pdf", "docx", "txt"])

if uploaded_file:
    with open(f"./contracts/{uploaded_file.name}", "wb") as f:
        f.write(uploaded_file.getbuffer())
    text = load_contract(f"./contracts/{uploaded_file.name}")
    clauses = split_into_clauses(text, max_len=600)

    st.success(f"✅ 成功載入合約，共 {len(clauses)} 個條款")

    tab1, tab2, tab3, tab4 = st.tabs(["📖 條款逐條分析", "📌 合約摘要", "⚠️ 風險重點", "⚖️ 法律檢索"])

    clause_analyses = []

    # 條款逐條分析
    with tab1:
        for i, clause in enumerate(clauses):
            with st.spinner(f"分析第 {i+1} 條款中..."):
                draft = generate_answer(clause, [clause])
                analysis = verify_answer(clause, draft, [clause])
                clause_analyses.append((clause, analysis))
                st.markdown(f"### 條款 {i+1}")
                st.info(clause)
                st.write(analysis)

    # 合約摘要
    with tab2:
        with st.spinner("正在生成合約摘要..."):
            draft_summary = generate_answer("請總結此合約", [text[:6000]])
            summary = verify_answer("請總結此合約", draft_summary, [text[:6000]])
        st.subheader("📌 合約摘要")
        st.write(summary)

    # 風險重點
    with tab3:
        with st.spinner("正在生成風險重點..."):
            draft_risks = generate_answer("請找出合約中的風險", [text[:6000]])
            risks = verify_answer("請找出合約中的風險", draft_risks, [text[:6000]])
        st.subheader("⚠️ 風險重點")
        st.write(risks)

    # 法律檢索
    with tab4:
        query = st.text_input("輸入法律問題（結合 RAG 檢索）")
        if query:
            candidates = hybrid_search(query, n=10)
            reranked = rerank(query, candidates, top_k=3)
            context_texts = [doc for (doc, _, _, _), _ in reranked]

            draft_answer = generate_answer(query, context_texts)
            final_answer = verify_answer(query, draft_answer, context_texts)

            st.subheader("🧠 法律回答")
            st.write(final_answer)

# 報告下載
def generate_word_report(summary, risks, clause_analyses):
    doc = Document()
    doc.add_heading("合約分析報告", level=1)
    doc.add_heading("📌 合約摘要", level=2)
    doc.add_paragraph(summary if summary else "無摘要")
    doc.add_heading("⚠️ 潛在風險", level=2)
    doc.add_paragraph(risks if risks else "未檢測到明顯風險。")
    doc.add_heading("📑 條款逐條分析", level=2)
    for i, (clause, analysis) in enumerate(clause_analyses, 1):
        doc.add_heading(f"條款 {i}", level=3)
        doc.add_paragraph(clause, style="Intense Quote")
        doc.add_paragraph(f"🔎 分析: {analysis}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def save_json_report(summary, risks, clause_analyses, filename="contract_analysis.json"):
    data = {
        "generated_at": datetime.now().isoformat(),
        "summary": summary,
        "risks": risks,
        "clauses": [{"clause": c, "analysis": a} for c, a in clause_analyses]
    }
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return filename

if uploaded_file:
    if st.button("📥 下載 Word 報告"):
        buf = generate_word_report(summary, risks, clause_analyses)
        st.download_button(
            label="⬇️ 下載合約分析報告 (Word)",
            data=buf,
            file_name="contract_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if st.button("📥 下載 JSON 報告"):
        json_file = save_json_report(summary, risks, clause_analyses)
        st.success(f"✅ JSON 記錄已保存到本地：{json_file}")
